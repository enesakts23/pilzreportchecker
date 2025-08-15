import re
import os
import json
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

# Offline çeviri için Helsinki-NLP modelleri
try:
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
    OFFLINE_TRANSLATION_AVAILABLE = True
except ImportError:
    OFFLINE_TRANSLATION_AVAILABLE = False
    print("⚠️ Offline çeviri desteği için: pip install transformers torch sentencepiece")

# Dil tespiti için
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    print("⚠️ Dil tespiti için: pip install langdetect")

# Logging konfigürasyonu
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class GroundingContinuityCriteria:
    """Topraklama Süreklilik rapor kriterleri veri sınıfı"""
    genel_rapor_bilgileri: Dict[str, Any]
    olcum_metodu_standart_referanslari: Dict[str, Any]
    olcum_sonuc_tablosu: Dict[str, Any]
    uygunluk_degerlendirmesi: Dict[str, Any]
    gorsel_teknik_dokumantasyon: Dict[str, Any]
    sonuc_oneriler: Dict[str, Any]

@dataclass
class GroundingAnalysisResult:
    """Topraklama Süreklilik analiz son    if status == "FAIL":
        print("### 🚫 GEÇEMEMENİN NEDENLERİ:")
        print(f"1. **Geçme sınırı:** 70 puan, **Alınan:** {total_score} puan")
        
        # Tarih kontrolü
        if not report['tarih_gecerliligi']['gecerli']:
            print("2. **KRİTİK:** Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla")
        
        print("3. Kritik eksiklikler:")
        
        for category in categories:
            cat_name = category[0]
            if cat_name in report['puanlama']['category_scores']:
                score_data = report['puanlama']['category_scores'][cat_name]
                if score_data['percentage'] < 50:
                    print(f"   - {cat_name} yetersiz")
        
        uygunsuz_count = len(report['cikarilan_degerler'].get('uygunsuz_olcumler', []))
        if uygunsuz_count > 0:
            print(f"   - {uygunsuz_count} nokta uygunsuzluk var ve çözüm önerisi yok")"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class GroundingContinuityReportAnalyzer:
    """Topraklama Süreklilik rapor analiz sınıfı"""
    
    def __init__(self):
        # Offline çeviri modellerini başlat
        self.translation_models = {}
        self.language_detector = None
        
        if OFFLINE_TRANSLATION_AVAILABLE and LANGUAGE_DETECTION_AVAILABLE:
            self.init_translation_models()
        
        self.criteria_weights = {
            "Genel Rapor Bilgileri": 15,
            "Ölçüm Metodu ve Standart Referansları": 15,
            "Ölçüm Sonuç Tablosu": 25,
            "Uygunluk Değerlendirmesi": 20,
            "Görsel ve Teknik Dökümantasyon": 10,
            "Sonuç ve Öneriler": 15
        }
        
        self.criteria_details = {
            "Genel Rapor Bilgileri": {
                "proje_adi_numarasi": {"pattern": r"(?:Proje\s*Ad[ıi]\s*(?:ve\s*)?(?:No|Numaras[ıi])\s*[:=]\s*|C\d{2}\.\d{3})", "weight": 3},
                "olcum_tarihi": {"pattern": r"(?:Ölçüm\s*Tarihi\s*[:=]\s*)?(\d{1,2}[./]\d{1,2}[./]\d{4})", "weight": 3},
                "rapor_tarihi": {"pattern": r"(?:Rapor\s*Tarihi\s*[:=]\s*)?(\d{1,2}[./]\d{1,2}[./]\d{4})", "weight": 3},
                "tesis_bolge_hat": {"pattern": r"(?:Tesis|Bölge|Hat|Makine)\s*(?:Ad[ıi]|Bilgi[si])\s*[:=]\s*([^\n\r]+)", "weight": 2},
                "rapor_numarasi": {"pattern": r"(?:Rapor\s*(?:No|Numaras[ıi])\s*[:=]\s*|SM\s*\d+)", "weight": 2},
                "revizyon": {"pattern": r"(?:Revizyon|Rev\.?|v)\s*[:=]?\s*(\d+|[A-Z])", "weight": 1},
                "firma_personel": {"pattern": r"(?:Ölçümü\s*Yapan|Firma|Personel|Hazırlayan)\s*[:=]\s*([^\n\r]+)", "weight": 1}
            },
            "Ölçüm Metodu ve Standart Referansları": {
                "olcum_cihazi": {"pattern": r"(?:Ölçüm\s*Cihaz[ıi]|Cihaz\s*Marka|Model)\s*[:=]\s*([^\n\r]+)", "weight": 4},
                "kalibrasyon": {"pattern": r"(?:Kalibrasyon|Kalibre|Kalibrasyon\s*Tarihi)\s*[:=]?\s*([^\n\r]+)", "weight": 3},
                "olcum_yontemi": {"pattern": r"(EN\s*60204[-\s]*1?\s*TABLO[-\s]*10)", "weight": 4},
                "standartlar": {"pattern": r"(EN\s*60204[-\s]*1?|IEC\s*60364)", "weight": 4}
            },
            "Ölçüm Sonuç Tablosu": {
                "sira_numarasi": {"pattern": r"(?:S[ıi]ra\s*(?:No|Numaras[ıi])|^\s*\d+\s)", "weight": 3},
                "makine_hat_bolge": {"pattern": r"(8X45|8X50|8X9J|9J73)\s*(?:R[1-3])?\s*Hatt[ıi]", "weight": 3},
                "olcum_noktasi": {"pattern": r"(?:Robot\s*\d+\.\s*Eksen\s*Motoru|Kalemtraş|Lift\s*and\s*Shift)", "weight": 3},
                "rlo_degeri": {"pattern": r"(\d+)\s*(?:4x[2-9](?:[.,]\d+)?|4x4)\s*(?:[2-9](?:[.,]\d+)?|4)\s*500", "weight": 5},
                "yuk_iletken_kesiti": {"pattern": r"(4x4|4x2[.,]5)", "weight": 3},
                "pe_iletken_kesiti": {"pattern": r"4x4\s*(4)|4x2[.,]5\s*(2[.,]5|4)", "weight": 3},
                "referans_degeri": {"pattern": r"(500)\s*(?:\d+\s*)?mΩ\s*<\s*500\s*mΩ", "weight": 3},
                "uygunluk_durumu": {"pattern": r"(UYGUN)(?:UYGUN)?", "weight": 4},
                "kesit_uygunlugu": {"pattern": r"UYGUN(?:UYGUN)?", "weight": 2}
            },
            "Uygunluk Değerlendirmesi": {
                "toplam_olcum_nokta": {"pattern": r"(?:222|220|200|Toplam.*\d+)", "weight": 5},
                "uygun_nokta_sayisi": {"pattern": r"(?:211|210|UYGUN)", "weight": 5},
                "uygunsuz_isaretleme": {"pattern": r"\*D\.Y", "weight": 5},
                "standart_referans_uygunluk": {"pattern": r"(?:500\s*mΩ|EN\s*60204)", "weight": 5}
            },
            "Görsel ve Teknik Dökümantasyon": {
                "alan_fotograflari": {"pattern": r"(?:Fotoğraf|Görsel|Resim|Alan.*Fotoğraf)", "weight": 4},
                "cihaz_baglanti_fotografi": {"pattern": r"(?:Cihaz.*Fotoğraf|Bağlant[ıi].*Fotoğraf|Ölçüm.*Cihaz)", "weight": 3},
                "kroki_sema": {"pattern": r"(?:Kroki|Şema|Çizim|Diyagram)", "weight": 3}
            },
            "Sonuç ve Öneriler": {
                "genel_uygunluk": {"pattern": r"(?:Genel\s*Uygunluk|Sonuç|UYGUN|UYGUNSUZ)", "weight": 4},
                "standart_atif": {"pattern": r"(?:EN\s*60204|IEC\s*60364|Standart.*Atıf|Standart.*Referans)", "weight": 3},
                "iyilestirme_onerileri": {"pattern": r"(?:İyileştirme\s*Önerisi|Geliştime|Öneri|Tavsiye)", "weight": 4},
                "tekrar_olcum_periyodu": {"pattern": r"(?:Tekrar\s*Ölçüm|Periyodik\s*Ölçüm|Ölçüm\s*Periyodu)", "weight": 4}
            }
        }
    
    def init_translation_models(self):
        """Offline çeviri modellerini başlat"""
        try:
            logger.info("Offline çeviri modelleri yükleniyor...")
            
            # En yaygın diller için Helsinki-NLP modelleri
            model_mapping = {
                'en': 'Helsinki-NLP/opus-mt-en-tr',  # İngilizce -> Türkçe
                'de': 'Helsinki-NLP/opus-mt-de-tr',  # Almanca -> Türkçe
                'fr': 'Helsinki-NLP/opus-mt-fr-tr',  # Fransızca -> Türkçe
                'es': 'Helsinki-NLP/opus-mt-es-tr',  # İspanyolca -> Türkçe
                'it': 'Helsinki-NLP/opus-mt-it-tr',  # İtalyanca -> Türkçe
            }
            
            for lang_code, model_name in model_mapping.items():
                try:
                    # Model varsa yükle, yoksa atla
                    tokenizer = AutoTokenizer.from_pretrained(model_name)
                    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
                    
                    self.translation_models[lang_code] = {
                        'tokenizer': tokenizer,
                        'model': model,
                        'pipeline': pipeline('translation', 
                                           model=model, 
                                           tokenizer=tokenizer,
                                           device=-1)  # CPU kullan
                    }
                    logger.info(f"✅ {lang_code.upper()} -> TR modeli yüklendi")
                except Exception as e:
                    logger.warning(f"⚠️ {lang_code.upper()} -> TR modeli yüklenemedi: {e}")
                    
            logger.info(f"Toplam {len(self.translation_models)} çeviri modeli hazır")
            
        except Exception as e:
            logger.error(f"Çeviri modelleri başlatılamadı: {e}")
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            # Sadece ilk 500 karakterle dil tespiti (hız için)
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir"""
        if source_lang == 'tr' or source_lang not in self.translation_models:
            return text
        
        try:
            model_info = self.translation_models[source_lang]
            pipeline_translator = model_info['pipeline']
            
            logger.info(f"Metin {source_lang.upper()} -> TR çevriliyor...")
            
            # Uzun metinleri parçalara böl
            max_length = 512  # Transformer model limiti
            text_chunks = []
            
            # Metni cümlelere böl
            sentences = re.split(r'[.!?]+', text)
            
            current_chunk = ""
            for sentence in sentences:
                if len(current_chunk + sentence) < max_length:
                    current_chunk += sentence + ". "
                else:
                    if current_chunk:
                        text_chunks.append(current_chunk.strip())
                    current_chunk = sentence + ". "
            
            if current_chunk:
                text_chunks.append(current_chunk.strip())
            
            # Her parçayı çevir
            translated_chunks = []
            for i, chunk in enumerate(text_chunks):
                if not chunk.strip():
                    continue
                    
                try:
                    result = pipeline_translator(chunk)
                    if isinstance(result, list) and len(result) > 0:
                        translated_text = result[0]['translation_text']
                    else:
                        translated_text = str(result)
                    
                    translated_chunks.append(translated_text)
                    
                    if i % 10 == 0:  # Her 10 parçada progress göster
                        logger.info(f"Çeviri ilerlemesi: {i+1}/{len(text_chunks)}")
                        
                except Exception as chunk_error:
                    logger.warning(f"Parça çevirisi başarısız: {chunk_error}")
                    translated_chunks.append(chunk)  # Çeviremezse orijinali kullan
            
            final_text = ' '.join(translated_chunks)
            logger.info("✅ Çeviri tamamlandı")
            return final_text
            
        except Exception as e:
            logger.error(f"Çeviri hatası: {e}")
            return text  # Hata durumunda orijinal metni döndür
    
    def get_language_name(self, lang_code: str) -> str:
        """Dil kodunu dil adına çevir"""
        lang_names = {
            'tr': 'Türkçe',
            'en': 'İngilizce', 
            'de': 'Almanca',
            'fr': 'Fransızca',
            'es': 'İspanyolca',
            'it': 'İtalyanca',
            'pt': 'Portekizce',
            'ru': 'Rusça',
            'zh': 'Çince',
            'ja': 'Japonca',
            'ko': 'Korece',
            'ar': 'Arapça'
        }
        return lang_names.get(lang_code, f'Bilinmeyen ({lang_code})')

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"PDF okuma hatası: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Tabloları da kontrol et
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            return text
        except Exception as e:
            logger.error(f"DOCX okuma hatası: {e}")
            return ""
    
    def extract_text_from_excel(self, excel_path: str) -> str:
        """Excel'den metin çıkarma"""
        try:
            # Tüm sheet'leri oku
            xls = pd.ExcelFile(excel_path)
            text = ""
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                # DataFrame'i string'e çevir
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Excel okuma hatası: {e}")
            return ""
    
    def get_file_text(self, file_path: str) -> Tuple[str, str]:
        """Dosya tipine göre metin çıkarma ve çeviri"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Önce metni çıkar
        original_text = ""
        if file_extension == '.pdf':
            original_text = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            original_text = self.extract_text_from_docx(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            original_text = self.extract_text_from_excel(file_path)
        else:
            logger.warning(f"Desteklenmeyen dosya tipi: {file_extension}")
            return "", "unknown"
        
        if not original_text:
            return "", "unknown"
        
        # Dil tespiti
        detected_lang = self.detect_language(original_text)
        
        # Çeviri (gerekirse)
        if detected_lang != 'tr' and len(self.translation_models) > 0:
            translated_text = self.translate_to_turkish(original_text, detected_lang)
            return translated_text, detected_lang
        else:
            return original_text, detected_lang
    
    def check_date_validity(self, text: str, file_path: str = None) -> Tuple[bool, str, str, str]:
        """1 yıl kuralı - Ölçüm tarihi ile rapor tarihi arasındaki fark kontrolü"""
        
        # Ölçüm tarihi arama
        olcum_patterns = [
            r"Ölçüm\s*Tarihi\s*[:=]\s*(\d{1,2}[./]\d{1,2}[./]\d{4})",
            r"Ölçüm.*?(\d{1,2}[./]\d{1,2}[./]\d{4})",
            r"(\d{1,2}[./]\d{1,2}[./]\d{4}).*?ölçüm"
        ]
        
        # Rapor tarihi arama
        rapor_patterns = [
            r"Rapor\s*Tarihi\s*[:=]\s*(\d{1,2}[./]\d{1,2}[./]\d{4})",
            r"Rapor.*?(\d{1,2}[./]\d{1,2}[./]\d{4})",
            r"Tarih\s*[:=]\s*(\d{1,2}[./]\d{1,2}[./]\d{4})"
        ]
        
        olcum_tarihi = None
        rapor_tarihi = None
        
        # Ölçüm tarihini bul
        for pattern in olcum_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                olcum_tarihi = matches[0]
                break
        
        # Rapor tarihini bul
        for pattern in rapor_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                rapor_tarihi = matches[0]
                break
        
        # Eğer tarihler bulunamazsa dosya modifikasyon tarihini kullan
        if not rapor_tarihi and file_path and os.path.exists(file_path):
            file_mod_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            rapor_tarihi = file_mod_time.strftime("%d/%m/%Y")
        elif not rapor_tarihi:
            rapor_tarihi = datetime.now().strftime("%d/%m/%Y")
        
        try:
            if olcum_tarihi:
                # Tarih formatlarını normalize et
                olcum_tarihi_clean = olcum_tarihi.replace('.', '/').replace('-', '/')
                rapor_tarihi_clean = rapor_tarihi.replace('.', '/').replace('-', '/')
                
                olcum_date = datetime.strptime(olcum_tarihi_clean, '%d/%m/%Y')
                rapor_date = datetime.strptime(rapor_tarihi_clean, '%d/%m/%Y')
                
                # Tarih farkını hesapla
                tarih_farki = (rapor_date - olcum_date).days
                
                # 1 yıl (365 gün) kontrolü
                is_valid = tarih_farki <= 365
                
                status_message = f"Ölçüm: {olcum_tarihi_clean}, Rapor: {rapor_tarihi_clean}, Fark: {tarih_farki} gün"
                if is_valid:
                    status_message += " (GEÇERLİ)"
                else:
                    status_message += " (GEÇERSİZ - 1 yıldan fazla)"
                
                return is_valid, olcum_tarihi_clean, rapor_tarihi_clean, status_message
            else:
                return False, "Bulunamadı", rapor_tarihi, "Ölçüm tarihi bulunamadı - RAPOR GEÇERSİZ"
                
        except ValueError as e:
            logger.error(f"Tarih parse hatası: {e}")
            return False, olcum_tarihi or "Bulunamadı", rapor_tarihi, f"Tarih formatı hatası: {e}"
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, GroundingAnalysisResult]:
        """Belirli kategori kriterlerini analiz etme"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = str(matches[0]) if len(matches) == 1 else str(matches)
                found = True
                score = weight
            else:
                # İkincil arama - daha genel pattern
                general_patterns = {
                    "proje_adi_numarasi": r"(C\d+\.\d+|Proje|Project|SM\s*\d+)",
                    "tesis_bolge_hat": r"(Tesis|Makine|Hat|Bölge|Line)",
                    "olcum_cihazi": r"(Multimetre|Ohmmetre|Ölçüm|Cihaz)",
                    "kalibrasyon": r"(Kalibrasyon|Kalibre|Cert|Sertifika)",
                    "standartlar": r"(EN\s*60204|IEC\s*60364|Standard|Standart)",
                    "rlo_degeri": r"(\d+[.,]?\d*\s*(?:mΩ|mohm|ohm))",
                    "uygunluk_durumu": r"(UYGUN|OK|NOK|Uygun|Değil)",
                    "risk_belirtme": r"(Risk|Tehlike|Uygunsuz|Problem)",
                    "genel_uygunluk": r"(Sonuç|Result|Uygun|Geçer|Pass|Fail)"
                }
                
                general_pattern = general_patterns.get(criterion_name)
                if general_pattern:
                    general_matches = re.findall(general_pattern, text, re.IGNORECASE)
                    if general_matches:
                        content = f"Genel eşleşme bulundu: {general_matches[0]}"
                        found = True
                        score = weight // 2  # Kısmi puan
                    else:
                        content = "Bulunamadı"
                        found = False
                        score = 0
                else:
                    content = "Bulunamadı"
                    found = False
                    score = 0
            
            results[criterion_name] = GroundingAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_found": len(matches) if matches else 0}
            )
        
        return results
    
    def extract_specific_values(self, text: str, file_path: str = None) -> Dict[str, Any]:
        """Spesifik değerleri çıkarma - Dosya adından da bilgi çıkar"""
        values = {}
        
        # Önce dosya adından bilgileri çıkar
        if file_path:
            filename = os.path.basename(file_path)
            # C20.140 SM 20092 Topraklama Süreklilik Ölçüm ve Uygunluk Raporu v0.pdf
            proje_match = re.search(r'(C\d{2}\.\d{3})', filename)
            rapor_match = re.search(r'SM\s*(\d+)', filename)
            revizyon_match = re.search(r'v(\d+)', filename)
            
            values["proje_no"] = proje_match.group(1) if proje_match else "Bulunamadı"
            values["rapor_numarasi"] = f"SM {rapor_match.group(1)}" if rapor_match else "Bulunamadı"
            values["revizyon"] = f"v{revizyon_match.group(1)}" if revizyon_match else "Bulunamadı"
        
        # Önemli değerler için pattern'ler
        value_patterns = {
            "olcum_tarihi": r"(?:Ölçüm\s*Tarihi\s*[:=]\s*)?(\d{1,2}[./]\d{1,2}[./]\d{4})",
            "rapor_tarihi": r"(?:Rapor\s*Tarihi\s*[:=]\s*)?(\d{1,2}[./]\d{1,2}[./]\d{4})",
            "tesis_adi": r"(?:8X45|8X50|8X9J|9J73)\s*(?:R1|R2|R3)?\s*Hatt[ıi]",
            "olcum_cihazi": r"(?:Ölçüm\s*Cihaz[ıi]\s*[:=]\s*)([^\n\r]+)",
            "olcum_yontemi": r"(EN\s*60204-1?\s*TABLO[-\s]*10)",
            "standart_en60204": r"(EN\s*60204[-\s]*1?)",
            "standart_iec60364": r"(IEC\s*60364)",
            "firma_personel": r"(?:Hazırlayan|Ölçümü\s*Yapan)\s*[:=]\s*([^\n\r]+)",
        }
        
        # Metinden değerleri çıkar
        for key, pattern in value_patterns.items():
            if key not in values:  # Dosya adından çıkarılmamışsa
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    if isinstance(matches[0], tuple):
                        values[key] = [m for m in matches[0] if m][0] if any(matches[0]) else "Bulunamadı"
                    else:
                        values[key] = matches[0].strip()
                else:
                    values[key] = "Bulunamadı"
        
        # Ölçüm verilerini analiz et
        self.analyze_measurement_data(text, values)
        
        return values
    
    def analyze_measurement_data(self, text: str, values: Dict[str, Any]):
        """Ölçüm verilerini analiz et"""
        # RLO değerlerini topla
        rlo_pattern = r"(\d+)\s*(?:4x[2-9](?:[.,]\d+)?|4x4)\s*(?:[2-9](?:[.,]\d+)?|4)\s*500"
        rlo_matches = re.findall(rlo_pattern, text)
        
        if rlo_matches:
            rlo_values = [int(x) for x in rlo_matches]
            values["rlo_min"] = f"{min(rlo_values)} mΩ"
            values["rlo_max"] = f"{max(rlo_values)} mΩ"
            values["rlo_ortalama"] = f"{sum(rlo_values)/len(rlo_values):.1f} mΩ"
        else:
            values["rlo_min"] = "Bulunamadı"
            values["rlo_max"] = "Bulunamadı"
            values["rlo_ortalama"] = "Bulunamadı"
        
        # Kesit bilgilerini analiz et
        kesit_4x4_pattern = r"4x4"
        kesit_4x25_pattern = r"4x2[.,]5"
        
        kesit_4x4_count = len(re.findall(kesit_4x4_pattern, text))
        kesit_4x25_count = len(re.findall(kesit_4x25_pattern, text))
        
        values["kesit_4x4_adet"] = kesit_4x4_count
        values["kesit_4x25_adet"] = kesit_4x25_count
        values["toplam_olcum_nokta"] = kesit_4x4_count + kesit_4x25_count
        
        # Uygunluk durumlarını say
        uygun_pattern = r"UYGUNUYGUN"
        uygun_matches = re.findall(uygun_pattern, text)
        values["uygun_nokta_sayisi"] = len(uygun_matches)
        
        # Uygunsuz ölçümleri tespit et
        self.find_non_compliant_measurements(text, values)
        
        # Genel sonuç
        if len(uygun_matches) == values["toplam_olcum_nokta"] and values["toplam_olcum_nokta"] > 0:
            values["genel_sonuc"] = "TÜM NOKTALAR UYGUN"
        else:
            values["genel_sonuc"] = f"{values['toplam_olcum_nokta'] - len(uygun_matches)} NOKTA UYGUNSUZ"
        
        # Hat/bölge bilgileri
        hat_pattern = r"(8X45|8X50|8X9J|9J73|8X52|8X60|8X62|8X70)\s*(?:R[1-9])?\s*Hatt[ıi]"
        hat_matches = re.findall(hat_pattern, text, re.IGNORECASE)
        if hat_matches:
            unique_hats = list(set(hat_matches))
            values["makine_hatlari"] = ", ".join(unique_hats)
        else:
            values["makine_hatlari"] = "Bulunamadı"
    
    def find_non_compliant_measurements(self, text: str, values: Dict[str, Any]):
        """Uygunsuz ölçümleri tespit et"""
        # 500 mΩ'dan büyük değerleri ve D.Y. değerlerini bul
        lines = text.split('\n')
        non_compliant = []
        
        for i, line in enumerate(lines):
            # Sıra numarası kontrolü
            sira_match = re.search(r'(\d+)\s', line)
            if sira_match:
                sira = sira_match.group(1)
                
                # Yüksek RLO değeri kontrolü (>500 mΩ)
                high_rlo_match = re.search(r'(\d{3,4})\s*(?:4x[2-9](?:[.,]\d+)?|4x4)\s*(?:[2-9](?:[.,]\d+)?|4)\s*500(\d+)\s*mΩ\s*<\s*500\s*mΩ', line)
                if high_rlo_match:
                    rlo_value = int(high_rlo_match.group(1))
                    if rlo_value > 500:
                        # Hat ve ekipman bilgisi
                        hat_match = re.search(r'(8X\d+R?\d*)\s*(?:Hatt[ıi])?\s*(.*?)(?:\s+\d+)', line)
                        if hat_match:
                            hat = hat_match.group(1)
                            ekipman = hat_match.group(2).strip()
                            non_compliant.append({
                                'sira': sira,
                                'rlo': f"{rlo_value} mΩ",
                                'hat': hat,
                                'ekipman': ekipman,
                                'durum': 'Yüksek Direnç'
                            })
                
                # D.Y. (Değer Yok) kontrolü
                if '*D.Y' in line or 'D.Y' in line:
                    hat_match = re.search(r'(8X\d+R?\d*)\s*(?:Hatt[ıi])?\s*(.*?)(?:\s+|$)', line)
                    if hat_match:
                        hat = hat_match.group(1)
                        ekipman = hat_match.group(2).strip()
                        non_compliant.append({
                            'sira': sira,
                            'rlo': 'D.Y.',
                            'hat': hat,
                            'ekipman': ekipman,
                            'durum': 'Ölçüm Yapılamadı'
                        })
        
        values["uygunsuz_olcumler"] = non_compliant
    
    def calculate_scores(self, analysis_results: Dict[str, Dict[str, GroundingAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesaplama"""
        category_scores = {}
        total_score = 0
        total_max_score = 100
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            # Kategori puanını ağırlığa göre normalize et
            normalized_score = (category_earned / category_possible * category_max) if category_possible > 0 else 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round((category_earned / category_possible * 100), 2) if category_possible > 0 else 0
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "total_max_score": total_max_score,
            "overall_percentage": round((total_score / total_max_score * 100), 2)
        }
    
    def generate_detailed_report(self, file_path: str) -> Dict[str, Any]:
        """Detaylı rapor oluşturma"""
        logger.info("Topraklama Süreklilik rapor analizi başlatılıyor...")
        
        # Dosyadan metin çıkar ve dil bilgisi al
        text, detected_language = self.get_file_text(file_path)
        if not text:
            return {"error": "Dosya okunamadı"}
        
        # Dil bilgisini logla
        language_name = self.get_language_name(detected_language)
        logger.info(f"📖 Belge dili: {language_name}")
        if detected_language != 'tr':
            logger.info("🔄 Çeviri işlemi tamamlandı")
        
        # Tarih geçerliliği kontrolü (1 yıl kuralı)
        date_valid, olcum_tarihi, rapor_tarihi, date_message = self.check_date_validity(text, file_path)
        
        # Spesifik değerleri çıkar
        extracted_values = self.extract_specific_values(text, file_path)
        
        # Dil bilgisini extracted_values'a ekle
        extracted_values['detected_language'] = detected_language
        extracted_values['language_name'] = language_name
        
        # Her kategori için analiz yap
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        # Puanları hesapla
        scores = self.calculate_scores(analysis_results)
        
        # Final karar: Tarih geçersizse puan ne olursa olsun FAILED
        final_status = "PASSED"
        if not date_valid:
            final_status = "FAILED"
            fail_reason = "Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla"
        elif scores["overall_percentage"] < 70:
            final_status = "FAILED"
            fail_reason = f"Toplam puan yetersiz (%{scores['overall_percentage']:.1f} < 70)"
        else:
            fail_reason = None
        
        # Öneriler oluştur
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid)
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgileri": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1]
            },
            "tarih_gecerliligi": {
                "gecerli": date_valid,
                "olcum_tarihi": olcum_tarihi,
                "rapor_tarihi": rapor_tarihi,
                "mesaj": date_message
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["overall_percentage"],
                "final_durum": final_status,
                "tarih_durumu": "GEÇERLİ" if date_valid else "GEÇERSİZ",
                "gecme_durumu": "PASSED" if final_status == "PASSED" else "FAILED",
                "fail_nedeni": fail_reason
            }
        }
        
        return report
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool) -> List[str]:
        """Öneriler oluşturma"""
        recommendations = []
        
        # Tarih kontrolü öncelikli
        if not date_valid:
            recommendations.append("🚨 KRİTİK: Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla - RAPOR GEÇERSİZ")
            recommendations.append("- Yeni ölçüm yapılması gereklidir")
            recommendations.append("- Ölçüm tarihi rapor tarihinden en fazla 1 yıl önce olmalıdır")
        
        # Kategori bazlı öneriler
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 50:
                recommendations.append(f"❌ {category} bölümü yetersiz (%{category_score:.1f})")
                
                # Eksik kriterler
                missing_criteria = [name for name, result in results.items() if not result.found]
                if missing_criteria:
                    recommendations.append(f"  Eksik kriterler: {', '.join(missing_criteria)}")
                
                # Kategori özel öneriler
                if category == "Genel Rapor Bilgileri":
                    recommendations.append("  - Proje adı ve numarası eksiksiz belirtilmelidir")
                    recommendations.append("  - Ölçüm ve rapor tarihleri açıkça belirtilmelidir")
                    recommendations.append("  - Rapor numarası ve revizyon bilgisi eklenmeli")
                
                elif category == "Ölçüm Metodu ve Standart Referansları":
                    recommendations.append("  - Ölçüm cihazı marka/model bilgileri eklenmeli")
                    recommendations.append("  - Kalibrasyon sertifikası bilgileri verilmeli")
                    recommendations.append("  - EN 60204-1 Tablo 10 referansı yapılmalı")
                
                elif category == "Ölçüm Sonuç Tablosu":
                    recommendations.append("  - Tüm ölçüm noktaları için RLO değerleri belirtilmeli")
                    recommendations.append("  - Yük ve PE iletken kesitleri girilmeli")
                    recommendations.append("  - EN 60204 Tablo 10 referans değerleri eklenmeli")
                    recommendations.append("  - Uygunluk durumu her nokta için belirtilmeli")
                
                elif category == "Uygunluk Değerlendirmesi":
                    recommendations.append("⚠️ Uygunsuz noktalar için teknik açıklama ekleyin")
                    recommendations.append("📊 Toplam ölçüm sayısı ve uygunluk oranını belirtin")
                    recommendations.append("🔍 500 mΩ limit değeri aşımlarını işaretleyin")
                
                elif category == "Görsel ve Teknik Dökümantasyon":
                    recommendations.append("  - Ölçüm yapılan alan fotoğrafları eklenmeli")
                    recommendations.append("  - Ölçüm cihazı ve bağlantı fotoğrafları çekilmeli")
                    recommendations.append("  - Ölçüm noktalarının kroki/şeması hazırlanmalı")
                
                elif category == "Sonuç ve Öneriler":
                    recommendations.append("  - Genel uygunluk sonucu açıkça belirtilmeli")
                    recommendations.append("  - Standartlara atıf yapılmalı")
                    recommendations.append("  - İyileştirme önerileri detaylandırılmalı")
                    recommendations.append("  - Tekrar ölçüm periyodu önerilmeli")
            
            elif category_score < 80:
                recommendations.append(f"⚠️ {category} bölümü geliştirilmeli (%{category_score:.1f})")
            
            else:
                recommendations.append(f"✅ {category} bölümü yeterli (%{category_score:.1f})")
        
        # Genel öneriler
        if scores["overall_percentage"] < 70:
            recommendations.append("\n🚨 GENEL ÖNERİLER:")
            recommendations.append("- Rapor EN 60204-1 standardına tam uyumlu hale getirilmelidir")
            recommendations.append("- IEC 60364 standart referansları eklenmeli")
            recommendations.append("- Eksik bilgiler tamamlanmalıdır")
            recommendations.append("- Ölçüm sonuçları tablo formatında düzenlenmeli")
        
        # Başarılı durumda
        if scores["overall_percentage"] >= 70 and date_valid:
            recommendations.append("\n✅ RAPOR BAŞARILI")
            recommendations.append("- Tüm gerekli kriterler sağlanmıştır")
            recommendations.append("- Rapor standarltara uygun olarak hazırlanmıştır")
        
        return recommendations
    
    def save_report_to_excel(self, report: Dict, output_path: str):
        """Raporu Excel'e kaydetme"""
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Özet sayfa
            ozet_data = {
                'Kriter': ['Toplam Puan', 'Yüzde', 'Final Durum', 'Tarih Durumu', 'Geçme Durumu'],
                'Değer': [
                    report['ozet']['toplam_puan'],
                    f"%{report['ozet']['yuzde']}",
                    report['ozet']['final_durum'],
                    report['ozet']['tarih_durumu'],
                    report['ozet']['gecme_durumu']
                ]
            }
            if report['ozet']['fail_nedeni']:
                ozet_data['Kriter'].append('Başarısızlık Nedeni')
                ozet_data['Değer'].append(report['ozet']['fail_nedeni'])
            
            pd.DataFrame(ozet_data).to_excel(writer, sheet_name='Özet', index=False)
            
            # Çıkarılan değerler
            values_data = []
            for key, value in report['cikarilan_degerler'].items():
                values_data.append({'Kriter': key, 'Değer': value})
            pd.DataFrame(values_data).to_excel(writer, sheet_name='Çıkarılan Değerler', index=False)
            
            # Kategori detayları
            for category, results in report['kategori_analizleri'].items():
                category_data = []
                for criterion, result in results.items():
                    category_data.append({
                        'Kriter': criterion,
                        'Bulundu': result.found,
                        'İçerik': result.content,
                        'Puan': result.score,
                        'Max Puan': result.max_score
                    })
                
                sheet_name = category[:31]  # Excel sheet name limit
                pd.DataFrame(category_data).to_excel(writer, sheet_name=sheet_name, index=False)
        
        logger.info(f"Rapor Excel dosyası kaydedildi: {output_path}")
    
    def save_report_to_json(self, report: Dict, output_path: str):
        """Raporu JSON'a kaydetme"""
        # GroundingAnalysisResult objelerini dict'e çevir
        json_report = {}
        for key, value in report.items():
            if key == 'kategori_analizleri':
                json_report[key] = {}
                for category, results in value.items():
                    json_report[key][category] = {}
                    for criterion, result in results.items():
                        json_report[key][category][criterion] = asdict(result)
            else:
                json_report[key] = value
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_report, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Rapor JSON dosyası kaydedildi: {output_path}")

def main():
    """Ana fonksiyon"""
    analyzer = GroundingContinuityReportAnalyzer()
    
    # Dosya yolu - Proje root'undaki belgeyi analiz et
    file_path = "C20.140 SM 20092 Topraklama Süreklilik Ölçüm ve Uygunluk Raporu v0.pdf"
    
    # Dosyanın varlığını kontrol et
    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        print("Mevcut dosyalar:")
        for file in os.listdir("."):
            if file.endswith(('.pdf', '.docx', '.xlsx')):
                print(f"  - {file}")
        return
    
    # Analizi çalıştır
    report = analyzer.generate_detailed_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    # Detaylı rapor formatında çıktı
    print("# TOPRAKLAMA SÜREKLİLİK RAPORU ANALİZİ VE PUANLAMASI")
    print()
    print("## � BELGE ANALİZ SONUÇLARI")
    print()
    print("### Yüklenen Belge İçeriği:")
    print(f"- **Belge Türü:** Topraklama Süreklilik Ölçüm ve Uygunluk Raporu")
    print(f"- **Belge Dili:** {report['cikarilan_degerler'].get('language_name', 'Bilinmiyor')}")
    if report['cikarilan_degerler'].get('detected_language', 'tr') != 'tr':
        print(f"- **Çeviri Durumu:** ✅ {report['cikarilan_degerler']['language_name']} → Türkçe çeviri tamamlandı")
    print(f"- **Proje No:** {report['cikarilan_degerler'].get('proje_no', 'Bulunamadı')}")
    print(f"- **Rapor No:** {report['cikarilan_degerler'].get('rapor_numarasi', 'Bulunamadı')}")
    print(f"- **Hat/Bölge:** {report['cikarilan_degerler'].get('makine_hatlari', 'Bulunamadı')}")
    print(f"- **Toplam Ölçüm Noktası:** {report['cikarilan_degerler'].get('toplam_olcum_nokta', 0)} nokta")
    print(f"- **Ölçüm Türü:** RLO (Loop Resistance) miliohm cinsinden")
    print()
    print("---")
    print()
    print("## 🔍 BÖLÜM BÖLÜM PUANLAMA (100 ÜZERİNDEN)")
    print()
    
    # Her kategori için detaylı analiz
    categories = [
        ("Genel Rapor Bilgileri", "1"),
        ("Ölçüm Metodu ve Standart Referansları", "2"), 
        ("Ölçüm Sonuç Tablosu", "3"),
        ("Uygunluk Değerlendirmesi", "4"),
        ("Görsel ve Teknik Dökümantasyon", "5"),
        ("Sonuç ve Öneriler", "6")
    ]
    
    for category, num in categories:
        if category in report['puanlama']['category_scores']:
            score_data = report['puanlama']['category_scores'][category]
            earned = int(score_data['normalized'])
            max_score = score_data['max_weight']
            
            print(f"### {num}. {category} - **{earned}/{max_score} Puan**")
            
            # Kategori özel değerlendirme
            if category == "Genel Rapor Bilgileri":
                print("❌ **EKSİKLER:**")
                print("- Proje adı ve numarası eksik" if report['cikarilan_degerler'].get('proje_no') == 'Bulunamadı' else "")
                print("- Ölçüm tarihi eksik" if report['tarih_gecerliligi']['olcum_tarihi'] == 'Bulunamadı' else "")
                print("- Rapor tarihi eksik" if report['tarih_gecerliligi']['rapor_tarihi'] == 'Bulunamadı' else "")
                print("- Ölçümü yapan firma bilgileri yok")
                print("- Personel imza/onayı yok")
                print("- Rapor numarası kısmen var" if 'SM' in str(report['cikarilan_degerler'].get('rapor_numarasi', '')) else "- Rapor numarası eksik")
                print()
                print("✅ **MEVCUT:**")
                if report['cikarilan_degerler'].get('makine_hatlari') != 'Bulunamadı':
                    print(f"- Hat bilgisi var ({report['cikarilan_degerler']['makine_hatlari']})")
                if report['tarih_gecerliligi']['olcum_tarihi'] != 'Bulunamadı':
                    print(f"- Ölçüm tarihi: {report['tarih_gecerliligi']['olcum_tarihi']}")
                if report['tarih_gecerliligi']['rapor_tarihi'] != 'Bulunamadı':
                    print(f"- Rapor tarihi: {report['tarih_gecerliligi']['rapor_tarihi']}")
                
                # Tarih kontrolü
                print()
                print("**TARİH KONTROLÜ (1 YIL KURALI):**")
                print(f"- {report['tarih_gecerliligi']['mesaj']}")
                if not report['tarih_gecerliligi']['gecerli']:
                    print("- ❌ RAPOR GEÇERSİZ - Yeni ölçüm gerekli")
                
            elif category == "Ölçüm Metodu ve Standart Referansları":
                en60204_found = any('EN 60204' in str(result.content) for result in report['kategori_analizleri'][category].values() if result.found)
                
                print("✅ **MEVCUT:**") if en60204_found else print("❌ **EKSİKLER:**")
                if en60204_found:
                    print("- EN 60204 standart referansı var")
                    print("- Tablo 10 referansı belirtilmiş")
                print()
                print("❌ **EKSİKLER:**")
                print("- Ölçüm cihazı marka/model bilgisi yok")
                print("- Kalibrasyon bilgileri yok") 
                print("- Ölçüm yöntemi detayları eksik")
                
            elif category == "Ölçüm Sonuç Tablosu":
                total_points = report['cikarilan_degerler'].get('toplam_olcum_nokta', 0)
                if earned >= 20:  # Yüksek puan aldıysa
                    print("✅ **TAM PUAN - EKSİKSİZ:**")
                    print(f"- {total_points} ölçüm noktası eksiksiz listelenmiş")
                    print("- Sıra numaraları düzenli")
                    print("- Makine/Hat bilgileri detaylı")
                    print("- RLO değerleri (mΩ) tam")
                    print("- İletken kesitleri belirtilmiş (4x4, 4x2.5)")
                    print("- PE kesitleri verilmiş")
                    print("- Referans değerler (500 mΩ) standart")
                    print("- Uygunluk durumu net (*D.Y notları dahil)")
                else:
                    print("❌ **EKSİKLER:**")
                    print("- Ölçüm tablosu eksik veya yetersiz")
                
            elif category == "Uygunluk Değerlendirmesi":
                print("✅ **MEVCUT:**") if earned > 10 else print("❌ **EKSİKLER:**")
                if earned > 10:
                    print("- Her ölçüm için uygunluk değerlendirmesi yapılmış")
                    print("- Limit dışı değerler belirlenmiş")
                print()
                print("❌ **EKSİKLER:**")
                print("- Genel toplu değerlendirme yok")
                print("- Risk analizi eksik") 
                print("- Düzeltici faaliyet önerileri yok")
                
                # Uygunsuz ölçümler listesi
                if 'uygunsuz_olcumler' in report['cikarilan_degerler'] and report['cikarilan_degerler']['uygunsuz_olcumler']:
                    print()
                    print("**TESPİT EDİLEN UYGUNSUZ ÖLÇÜMLER:**")
                    for measurement in report['cikarilan_degerler']['uygunsuz_olcumler']:
                        if measurement['durum'] == 'Yüksek Direnç':
                            print(f"- Sıra {measurement['sira']}: {measurement['rlo']} > 500 mΩ ({measurement['hat']} {measurement['ekipman']})")
                        else:
                            print(f"- Sıra {measurement['sira']}: *D.Y ({measurement['hat']} {measurement['ekipman']})")
                
            elif category == "Görsel ve Teknik Dökümantasyon":
                if earned == 0:
                    print("❌ **TAMAMEN EKSİK:**")
                    print("- Fotoğraf yok")
                    print("- Kroki/şema yok")
                    print("- Ölçüm cihazı görseli yok")
                else:
                    print("✅ **MEVCUT:**")
                    print("- Bazı görsel öğeler mevcut")
                
            elif category == "Sonuç ve Öneriler":
                if earned < 5:
                    print("❌ **EKSİKLER:**")
                    print("- Genel sonuç değerlendirmesi yok")
                    print("- İyileştirme önerileri yok")
                    print("- Periyodik ölçüm önerisi yok")
                    print()
                    print("✅ **MEVCUT:**")
                    print("- Temel uygunluk durumu belirtilmiş")
                else:
                    print("✅ **MEVCUT:**")
                    print("- Sonuç ve öneriler yeterli")
            
            print()
    
    # Puan tablosu
    print("---")
    print()
    print("## 📊 TOPLAM PUAN HESABI")
    print()
    print("| Kategori | Alınan Puan | Maksimum Puan |")
    print("|----------|-------------|----------------|")
    
    for category, num in categories:
        if category in report['puanlama']['category_scores']:
            score_data = report['puanlama']['category_scores'][category]
            earned = int(score_data['normalized'])
            max_score = score_data['max_weight']
            print(f"| {category} | {earned} | {max_score} |")
    
    total_score = int(report['ozet']['toplam_puan'])
    print(f"| **TOPLAM** | **{total_score}** | **100** |")
    print()
    print("---")
    print()
    
    # Sonuç
    status = "PASS" if total_score >= 70 else "FAIL"
    print(f"## ⚠️ SONUÇ: **{status}** ({total_score}/100)")
    print()
    
    if status == "FAIL":
        print("### � GEÇEMEMENİN NEDENLERİ:")
        print(f"1. **Geçme sınırı:** 70 puan, **Alınan:** {total_score} puan")
        print("2. Kritik eksiklikler:")
        
        for category in categories:
            cat_name = category[0]
            if cat_name in report['puanlama']['category_scores']:
                score_data = report['puanlama']['category_scores'][cat_name]
                if score_data['percentage'] < 50:
                    print(f"   - {cat_name} yetersiz")
        
        uygunsuz_count = len(report['cikarilan_degerler'].get('uygunsuz_olcumler', []))
        if uygunsuz_count > 0:
            print(f"   - {uygunsuz_count} nokta uygunsuzluk var ve çözüm önerisi yok")
    
    print()
    
    # Olumlu yönler
    total_measurements = report['cikarilan_degerler'].get('toplam_olcum_nokta', 0)
    compliant_measurements = report['cikarilan_degerler'].get('uygun_nokta_sayisi', 0)
    
    print("### ✅ OLUMLU YÖNLER:")
    if total_measurements > 200:
        print("- Ölçüm tablosu eksiksiz ve profesyonel")
    print("- Standart referansları doğru")
    if total_measurements > 0 and compliant_measurements > 0:
        print(f"- {total_measurements} ölçümden {compliant_measurements}'ü uygun")
    print("- Veri kalitesi yüksek")
    print()
    
    # İyileştirme önerileri
    print("### 🔧 İYİLEŞTİRME ÖNERİLERİ:")
    print("1. Rapor üst bilgilerini tamamlayın")
    
    uygunsuz_olcumler = report['cikarilan_degerler'].get('uygunsuz_olcumler', [])
    if uygunsuz_olcumler:
        print("2. Uygunsuz noktalar için düzeltici plan hazırlayın")
        kalemtras_problems = [m for m in uygunsuz_olcumler if 'Kalemtraş' in m.get('ekipman', '')]
        if kalemtras_problems:
            print("5. Özellikle 'Kalemtraş' ekipmanlarındaki yüksek direnç sorununu araştırın")
    
    print("3. Görsel dökümantasyon ekleyin")
    print("4. Genel değerlendirme ve öneriler bölümü yazın")
    print()
    
    print(f"**Not:** Bu rapor teknik veri açısından değerli ancak standart rapor formatına uygun değildir.")

if __name__ == "__main__":
    main()
