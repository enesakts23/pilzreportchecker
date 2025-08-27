import re
import os
import json
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

# Offline çeviri için Helsinki-NLP modelleri
try:
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
    OFFLINE_TRANSLATION_AVAILABLE = True
except ImportError:
    OFFLINE_TRANSLATION_AVAILABLE = False
    print("⚠️ Offline çeviri desteği için: pip install transformers torch sentencepiece")

# Dil tespiti için
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    print("⚠️ Dil tespiti için: pip install langdetect")

# Logging konfigürasyonu
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class GroundingContinuityCriteria:
    """Topraklama Süreklilik rapor kriterleri veri sınıfı"""
    genel_rapor_bilgileri: Dict[str, Any]
    olcum_metodu_standart_referanslari: Dict[str, Any]
    olcum_sonuc_tablosu: Dict[str, Any]
    uygunluk_degerlendirmesi: Dict[str, Any]
    gorsel_teknik_dokumantasyon: Dict[str, Any]
    sonuc_oneriler: Dict[str, Any]

@dataclass
class GroundingAnalysisResult:
    """Topraklama Süreklilik analiz son    if status == "FAIL":
        print("### 🚫 GEÇEMEMENİN NEDENLERİ:")
        print(f"1. **Geçme sınırı:** 70 puan, **Alınan:** {total_score} puan")
        
        # Tarih kontrolü
        if not report['tarih_gecerliligi']['gecerli']:
            print("2. **KRİTİK:** Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla")
        
        print("3. Kritik eksiklikler:")
        
        for category in categories:
            cat_name = category[0]
            if cat_name in report['puanlama']['category_scores']:
                score_data = report['puanlama']['category_scores'][cat_name]
                if score_data['percentage'] < 50:
                    print(f"   - {cat_name} yetersiz")
        
        uygunsuz_count = len(report['cikarilan_degerler'].get('uygunsuz_olcumler', []))
        if uygunsuz_count > 0:
            print(f"   - {uygunsuz_count} nokta uygunsuzluk var ve çözüm önerisi yok")"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class GroundingContinuityReportAnalyzer:
    """Topraklama Süreklilik rapor analiz sınıfı"""
    
    def __init__(self):
        # Offline çeviri modellerini başlat
        self.translation_models = {}
        self.language_detector = None
        
        if OFFLINE_TRANSLATION_AVAILABLE and LANGUAGE_DETECTION_AVAILABLE:
            self.init_translation_models()
        
        self.criteria_weights = {
            "Genel Rapor Bilgileri": 15,
            "Ölçüm Metodu ve Standart Referansları": 15,
            "Ölçüm Sonuç Tablosu": 25,
            "Uygunluk Değerlendirmesi": 20,
            "Görsel ve Teknik Dökümantasyon": 10,
            "Sonuç ve Öneriler": 15
        }
        
        self.criteria_details = {
            "Genel Rapor Bilgileri": {
                "proje_adi_numarasi": {"pattern": r"(?:Project\s*(?:Name|No)|Proje\s*(?:Ad[ıi]|No)|Report\s*Title|Document\s*Title|E\d{2}\.\d{3}|C\d{2}\.\d{3}|T\d{2,3}[-.]?\d{3,4})", "weight": 3},
                "olcum_tarihi": {"pattern": r"(?:Measurement\s*Date|Ölçüm\s*Tarihi|Test\s*Date|Date\s*of\s*(?:Test|Measurement)|Measured\s*on|Tested\s*on|\d{1,2}[./\-]\d{1,2}[./\-]\d{4})", "weight": 3},
                "rapor_tarihi": {"pattern": r"(?:Report\s*Date|Rapor\s*Tarihi|Issue\s*Date|Document\s*Date|Prepared\s*on|Created\s*on|Date|Tarih|\d{1,2}[./\-]\d{1,2}[./\-]\d{4})", "weight": 3},
                "tesis_bolge_hat": {"pattern": r"(?:Customer|Müşteri|Client|Facility|Tesis|Plant|Factory|Company|Firma|Toyota|DANONE|Ford|BOSCH)", "weight": 2},
                "rapor_numarasi": {"pattern": r"(?:Report\s*No|Rapor\s*No|Document\s*No|Belge\s*No|E\d{2}\.\d{3}|C\d{2}\.\d{3}|SM\s*\d+|MCC\d+)", "weight": 2},
                "revizyon": {"pattern": r"(?:Version|Revizyon|Rev\.?|v)\s*[:=]?\s*(\d+|[A-Z])", "weight": 1},
                "firma_personel": {"pattern": r"(?:Prepared\s*by|Hazırlayan|Performed\s*by|Ölçümü\s*Yapan|Consultant|Engineer|PILZ)", "weight": 1}
            },
            "Ölçüm Metodu ve Standart Referansları": {
                "olcum_cihazi": {"pattern": r"(?:Measuring\s*Instrument|Ölçüm\s*Cihaz[ıi]|Test\s*Equipment|Multimeter|Multimetre|Ohmmeter|Instrument|Equipment|Device|Tester|Fluke|Metrix|Chauvin|Megger|Hioki)", "weight": 6},
                "kalibrasyon": {"pattern": r"(?:Calibration|Kalibrasyon|Kalibre|Certificate|Sertifika|Cal\s*Date)", "weight": 4},
                "standartlar": {"pattern": r"(?:EN\s*60204[-\s]*1?|IEC\s*60364|Standard|Standart)", "weight": 5}
            },
            "Ölçüm Sonuç Tablosu": {
                "sira_numarasi": {"pattern": r"(?:S[ıi]ra\s*(?:No|Numaras[ıi])|^\s*\d+\s)", "weight": 3},
                "makine_hat_bolge": {"pattern": r"(?:8X45|8X50|8X9J|9J73|8X52|8X60|8X62|8X70)\s*(?:R[1-9])?\s*(?:Hatt[ıi]|Line|Zone|Bölge)", "weight": 3},
                "olcum_noktasi": {"pattern": r"(?:Robot\s*\d+\.\s*Eksen\s*Motoru|Kalemtraş|Lift\s*and\s*Shift|Motor|Ekipman|Equipment|Device)", "weight": 3},
                "rlo_degeri": {"pattern": r"(\d+[.,]?\d*)\s*(?:mΩ|mohm|ohm|Ω)", "weight": 5},
                "yuk_iletken_kesiti": {"pattern": r"(?:4x4|4x2[.,]5|4x6|4x10|Yük\s*İletken|Load\s*Conductor|PE\s*İletken|PE\s*Conductor)", "weight": 4},
                "referans_degeri": {"pattern": r"(?:500\s*mΩ|500\s*ohm|500\s*Ω|EN\s*60204|IEC\s*60364)", "weight": 3},
                "uygunluk_durumu": {"pattern": r"(?:UYGUN|OK|PASS|Compliant|Uygun)", "weight": 4},
                "kesit_uygunlugu": {"pattern": r"(?:UYGUN|OK|PASS|Compliant|Uygun)", "weight": 3}
            },
            "Uygunluk Değerlendirmesi": {
                "toplam_olcum_nokta": {"pattern": r"(?:222|220|200|Toplam.*\d+)", "weight": 5},
                "uygun_nokta_sayisi": {"pattern": r"(?:211|210|UYGUN)", "weight": 5},
                "uygunsuz_isaretleme": {"pattern": r"\*D\.Y", "weight": 5, "reverse_logic": True},  # Uygunsuzluk bulunmazsa tam puan
                "standart_referans_uygunluk": {"pattern": r"(?:500\s*mΩ|EN\s*60204)", "weight": 5}
            },
            "Görsel ve Teknik Dökümantasyon": {
                "cihaz_baglanti_fotografi": {"pattern": r"(?:Cihaz.*Fotoğraf|Bağlant[ıi].*Fotoğraf|Ölçüm.*Cihaz|Photo|Image|Figure|Resim|Görsel)", "weight": 10}
            },
            "Sonuç ve Öneriler": {
                "genel_uygunluk": {"pattern": r"(?:Genel\s*Uygunluk|Sonuç|UYGUN|UYGUNSUZ|Result|Conclusion|Compliant|Non-compliant)", "weight": 8},
                "standart_atif": {"pattern": r"(?:EN\s*60204|IEC\s*60364|Standart.*Atıf|Standart.*Referans|Standard.*Reference)", "weight": 7}
            }
        }
    
    def init_translation_models(self):
        """Offline çeviri modellerini başlat"""
        try:
            logger.info("Offline çeviri modelleri yükleniyor...")
            
            # En yaygın diller için Helsinki-NLP modelleri
            model_mapping = {
                'en': 'Helsinki-NLP/opus-mt-en-tr',  # İngilizce -> Türkçe
                'de': 'Helsinki-NLP/opus-mt-de-tr',  # Almanca -> Türkçe
                'fr': 'Helsinki-NLP/opus-mt-fr-tr',  # Fransızca -> Türkçe
                'es': 'Helsinki-NLP/opus-mt-es-tr',  # İspanyolca -> Türkçe
                'it': 'Helsinki-NLP/opus-mt-it-tr',  # İtalyanca -> Türkçe
            }
            
            for lang_code, model_name in model_mapping.items():
                try:
                    # Model varsa yükle, yoksa atla
                    tokenizer = AutoTokenizer.from_pretrained(model_name)
                    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
                    
                    self.translation_models[lang_code] = {
                        'tokenizer': tokenizer,
                        'model': model,
                        'pipeline': pipeline('translation', 
                                           model=model, 
                                           tokenizer=tokenizer,
                                           device=-1)  # CPU kullan
                    }
                    logger.info(f"✅ {lang_code.upper()} -> TR modeli yüklendi")
                except Exception as e:
                    logger.warning(f"⚠️ {lang_code.upper()} -> TR modeli yüklenemedi: {e}")
                    
            logger.info(f"Toplam {len(self.translation_models)} çeviri modeli hazır")
            
        except Exception as e:
            logger.error(f"Çeviri modelleri başlatılamadı: {e}")
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            # Sadece ilk 500 karakterle dil tespiti (hız için)
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir"""
        if source_lang == 'tr' or source_lang not in self.translation_models:
            return text
        
        try:
            model_info = self.translation_models[source_lang]
            pipeline_translator = model_info['pipeline']
            
            logger.info(f"Metin {source_lang.upper()} -> TR çevriliyor...")
            
            # Uzun metinleri parçalara böl
            max_length = 512  # Transformer model limiti
            text_chunks = []
            
            # Metni cümlelere böl
            sentences = re.split(r'[.!?]+', text)
            
            current_chunk = ""
            for sentence in sentences:
                if len(current_chunk + sentence) < max_length:
                    current_chunk += sentence + ". "
                else:
                    if current_chunk:
                        text_chunks.append(current_chunk.strip())
                    current_chunk = sentence + ". "
            
            if current_chunk:
                text_chunks.append(current_chunk.strip())
            
            # Her parçayı çevir
            translated_chunks = []
            for i, chunk in enumerate(text_chunks):
                if not chunk.strip():
                    continue
                    
                try:
                    result = pipeline_translator(chunk)
                    if isinstance(result, list) and len(result) > 0:
                        translated_text = result[0]['translation_text']
                    else:
                        translated_text = str(result)
                    
                    translated_chunks.append(translated_text)
                    
                    if i % 10 == 0:  # Her 10 parçada progress göster
                        logger.info(f"Çeviri ilerlemesi: {i+1}/{len(text_chunks)}")
                        
                except Exception as chunk_error:
                    logger.warning(f"Parça çevirisi başarısız: {chunk_error}")
                    translated_chunks.append(chunk)  # Çeviremezse orijinali kullan
            
            final_text = ' '.join(translated_chunks)
            logger.info("✅ Çeviri tamamlandı")
            return final_text
            
        except Exception as e:
            logger.error(f"Çeviri hatası: {e}")
            return text  # Hata durumunda orijinal metni döndür
    
    def get_language_name(self, lang_code: str) -> str:
        """Dil kodunu dil adına çevir"""
        lang_names = {
            'tr': 'Türkçe',
            'en': 'İngilizce', 
            'de': 'Almanca',
            'fr': 'Fransızca',
            'es': 'İspanyolca',
            'it': 'İtalyanca',
            'pt': 'Portekizce',
            'ru': 'Rusça',
            'zh': 'Çince',
            'ja': 'Japonca',
            'ko': 'Korece',
            'ar': 'Arapça'
        }
        return lang_names.get(lang_code, f'Bilinmeyen ({lang_code})')

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"PDF okuma hatası: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Tabloları da kontrol et
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            return text
        except Exception as e:
            logger.error(f"DOCX okuma hatası: {e}")
            return ""
    
    def extract_text_from_excel(self, excel_path: str) -> str:
        """Excel'den metin çıkarma"""
        try:
            # Tüm sheet'leri oku
            xls = pd.ExcelFile(excel_path)
            text = ""
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                # DataFrame'i string'e çevir
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Excel okuma hatası: {e}")
            return ""
    
    def get_file_text(self, file_path: str) -> Tuple[str, str]:
        """Dosya tipine göre metin çıkarma ve çeviri"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Önce metni çıkar
        original_text = ""
        if file_extension == '.pdf':
            original_text = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            original_text = self.extract_text_from_docx(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            original_text = self.extract_text_from_excel(file_path)
        else:
            logger.warning(f"Desteklenmeyen dosya tipi: {file_extension}")
            return "", "unknown"
        
        if not original_text:
            return "", "unknown"
        
        # Dil tespiti
        detected_lang = self.detect_language(original_text)
        
        # Çeviri (gerekirse)
        if detected_lang != 'tr' and len(self.translation_models) > 0:
            translated_text = self.translate_to_turkish(original_text, detected_lang)
            return translated_text, detected_lang
        else:
            return original_text, detected_lang
    
    def normalize_date_string(self, date_str: str) -> str:
        """Tarih string'ini DD/MM/YYYY formatına çevir"""
        if not date_str or date_str == "Bulunamadı":
            return date_str
            
        # Ay isimleri çeviri tablosu
        month_names = {
            # İngilizce ay isimleri
            'jan': '01', 'january': '01',
            'feb': '02', 'february': '02', 
            'mar': '03', 'march': '03',
            'apr': '04', 'april': '04',
            'may': '05',
            'jun': '06', 'june': '06',
            'jul': '07', 'july': '07',
            'aug': '08', 'august': '08',
            'sep': '09', 'september': '09',
            'oct': '10', 'october': '10',
            'nov': '11', 'november': '11',
            'dec': '12', 'december': '12',
            
            # Türkçe ay isimleri
            'ocak': '01',
            'şubat': '02', 'subat': '02',
            'mart': '03',
            'nisan': '04',
            'mayıs': '05', 'mayis': '05',
            'haziran': '06',
            'temmuz': '07',
            'ağustos': '08', 'agustos': '08',
            'eylül': '09', 'eylul': '09',
            'ekim': '10',
            'kasım': '11', 'kasim': '11',
            'aralık': '12', 'aralik': '12'
        }
        
        # Çeşitli tarih formatlarını normalize et
        date_str = date_str.strip()
        
        # DD/MM/YYYY veya DD.MM.YYYY veya DD-MM-YYYY formatları
        if re.match(r'\d{1,2}[./\-]\d{1,2}[./\-]\d{4}', date_str):
            return date_str.replace('.', '/').replace('-', '/')
        
        # YYYY/MM/DD formatı
        if re.match(r'\d{4}[./\-]\d{1,2}[./\-]\d{1,2}', date_str):
            parts = re.split(r'[./\-]', date_str)
            return f"{parts[2].zfill(2)}/{parts[1].zfill(2)}/{parts[0]}"
        
        # DD Month YYYY formatı (örn: "18 Apr 2023" veya "18 Nisan 2023")
        month_pattern = r'(\d{1,2})\s+([a-zA-ZğıüşçöĞIÜŞÇÖ]+)\s+(\d{4})'
        match = re.match(month_pattern, date_str, re.IGNORECASE)
        if match:
            day, month_name, year = match.groups()
            month_name_lower = month_name.lower()
            if month_name_lower in month_names:
                month_num = month_names[month_name_lower]
                return f"{day.zfill(2)}/{month_num}/{year}"
        
        # Eğer hiçbir format eşleşmezse orijinal string'i döndür
        return date_str.replace('.', '/').replace('-', '/')
    
    def check_date_validity(self, text: str, file_path: str = None) -> Tuple[bool, str, str, str]:
        """1 yıl kuralı - Ölçüm tarihi ile rapor tarihi arasındaki fark kontrolü"""
        
        # Ölçüm tarihi arama - çok kapsamlı pattern'lar
        olcum_patterns = [
            # Türkçe formatlar
            r"(?:Ölçüm\s*Tarihi|Test\s*Tarihi|Ölçüm\s*Yapıldığı\s*Tarih)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(?:Ölçüm|Test).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4}).*?(?:ölçüm|test)",
            
            # İngilizce formatlar
            r"(?:Measurement\s*Date|Test\s*Date|Date\s*of\s*(?:Test|Measurement))\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(?:Measured\s*on|Tested\s*on)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4}).*?(?:measurement|test)",
            
            # Genel formatlar
            r"(\d{4}[./\-]\d{1,2}[./\-]\d{1,2})",  # YYYY/MM/DD
            r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})",
            r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4})"
        ]
        
        # Rapor tarihi arama - çok kapsamlı pattern'lar
        rapor_patterns = [
            # Türkçe formatlar
            r"(?:Rapor\s*Tarihi|Belge\s*Tarihi|Hazırlanma\s*Tarihi)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(?:Rapor|Belge).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(?:Hazırlayan|Hazırlandı)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            
            # İngilizce formatlar
            r"(?:Report\s*Date|Document\s*Date|Issue\s*Date|Prepared\s*Date)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(?:Prepared\s*on|Issued\s*on|Created\s*on)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            
            # Genel formatlar
            r"(?:Date|Tarih)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
            r"(\d{4}[./\-]\d{1,2}[./\-]\d{1,2})",  # YYYY/MM/DD
            r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})",
            r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4})"
        ]
        
        olcum_tarihi = None
        rapor_tarihi = None
        
        # Ölçüm tarihini bul
        for pattern in olcum_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                olcum_tarihi = matches[0]
                break
        
        # Rapor tarihini bul
        for pattern in rapor_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                rapor_tarihi = matches[0]
                break
        
        # Eğer tarihler bulunamazsa dosya modifikasyon tarihini kullan
        if not rapor_tarihi and file_path and os.path.exists(file_path):
            file_mod_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            rapor_tarihi = file_mod_time.strftime("%d/%m/%Y")
        elif not rapor_tarihi:
            rapor_tarihi = datetime.now().strftime("%d/%m/%Y")
        
        try:
            if olcum_tarihi:
                # Tarih formatlarını normalize et ve ay isimlerini çevir
                olcum_tarihi_clean = self.normalize_date_string(olcum_tarihi)
                rapor_tarihi_clean = self.normalize_date_string(rapor_tarihi)
                
                olcum_date = datetime.strptime(olcum_tarihi_clean, '%d/%m/%Y')
                rapor_date = datetime.strptime(rapor_tarihi_clean, '%d/%m/%Y')
                
                # Tarih farkını hesapla
                tarih_farki = (rapor_date - olcum_date).days
                
                # 1 yıl (365 gün) kontrolü
                is_valid = tarih_farki <= 365
                
                status_message = f"Ölçüm: {olcum_tarihi_clean}, Rapor: {rapor_tarihi_clean}, Fark: {tarih_farki} gün"
                if is_valid:
                    status_message += " (GEÇERLİ)"
                else:
                    status_message += " (GEÇERSİZ - 1 yıldan fazla)"
                
                return is_valid, olcum_tarihi_clean, rapor_tarihi_clean, status_message
            else:
                return False, "Bulunamadı", rapor_tarihi, "Ölçüm tarihi bulunamadı - RAPOR GEÇERSİZ"
                
        except ValueError as e:
            logger.error(f"Tarih parse hatası: {e}")
            return False, olcum_tarihi or "Bulunamadı", rapor_tarihi, f"Tarih formatı hatası: {e}"
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, GroundingAnalysisResult]:
        """Belirli kategori kriterlerini analiz etme"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            reverse_logic = criterion_data.get("reverse_logic", False)
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                if reverse_logic:
                    # Uygunsuzluk bulundu - düşük puan
                    content = f"Uygunsuzluk tespit edildi: {str(matches[:3])}"
                    found = True
                    score = weight // 3  # Düşük puan
                else:
                    content = str(matches[0]) if len(matches) == 1 else str(matches)
                    found = True
                    score = weight
            else:
                if reverse_logic:
                    # Uygunsuzluk bulunamadı - tam puan (iyi bir şey)
                    content = "Uygunsuzluk bulunamadı - Tüm ölçümler uygun"
                    found = True
                    score = weight  # Tam puan
                else:
                    # İkincil arama - daha genel pattern
                    general_patterns = {
                        "proje_adi_numarasi": r"(C\d+\.\d+|Proje|Project|SM\s*\d+)",
                        "tesis_bolge_hat": r"(Tesis|Makine|Hat|Bölge|Line)",
                        "olcum_cihazi": r"(Multimetre|Ohmmetre|Ölçüm|Cihaz)",
                        "kalibrasyon": r"(Kalibrasyon|Kalibre|Cert|Sertifika)",
                        "standartlar": r"(EN\s*60204|IEC\s*60364|Standard|Standart)",
                        "rlo_degeri": r"(\d+[.,]?\d*\s*(?:mΩ|mohm|ohm))",
                        "uygunluk_durumu": r"(UYGUN|OK|NOK|Uygun|Değil)",
                        "risk_belirtme": r"(Risk|Tehlike|Uygunsuz|Problem)",
                        "genel_uygunluk": r"(Sonuç|Result|Uygun|Geçer|Pass|Fail)"
                    }
                    
                    general_pattern = general_patterns.get(criterion_name)
                    if general_pattern:
                        general_matches = re.findall(general_pattern, text, re.IGNORECASE)
                        if general_matches:
                            content = f"Genel eşleşme bulundu: {general_matches[0]}"
                            found = True
                            score = weight // 2  # Kısmi puan
                        else:
                            content = "Bulunamadı"
                            found = False
                            score = 0
                    else:
                        content = "Bulunamadı"
                        found = False
                        score = 0
            
            results[criterion_name] = GroundingAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_found": len(matches) if matches else 0}
            )
        
        return results
    
    def extract_specific_values(self, text: str, file_path: str = None) -> Dict[str, Any]:
        """Spesifik değerleri çıkarma - Dosya adından da bilgi çıkar"""
        values = {}
        
        # Önce dosya adından bilgileri çıkar
        if file_path:
            filename = os.path.basename(file_path)
            
            # Proje numarası pattern'leri - farklı formatlar için
            proje_patterns = [
                r'(C\d{2}\.\d{3})',  # C20.140 formatı
                r'(E\d{2}\.\d{3})',  # E21.207 formatı
                r'(T\d{2,3}[-\.]?\d{3,4})',  # T21-MCC1201 formatı
                r'(\d{4,6})',        # 20092 gibi sayı formatı
                r'([A-Z]\d{2,3}[.-]\d{3,4})'  # Genel format
            ]
            
            # Rapor numarası pattern'leri
            rapor_patterns = [
                r'SM\s*(\d+)',
                r'MCC(\d+)',
                r'Report\s*No[\s:]*([A-Z0-9.-]+)',
                r'Rapor[\s:]*([A-Z0-9.-]+)'
            ]
            
            # Müşteri/firma bilgisi
            musteri_patterns = [
                r'Toyota',
                r'DANONE',
                r'Ford',
                r'BOSCH',
                r'P&G'
            ]
            
            # Dosya adından proje no çıkar
            proje_no = "Bulunamadı"
            for pattern in proje_patterns:
                match = re.search(pattern, filename, re.IGNORECASE)
                if match:
                    proje_no = match.group(1)
                    break
            values["proje_no"] = proje_no
            
            # Dosya adından rapor numarası çıkar
            rapor_no = "Bulunamadı"
            for pattern in rapor_patterns:
                match = re.search(pattern, filename, re.IGNORECASE)
                if match:
                    rapor_no = match.group(1)
                    break
            values["rapor_numarasi"] = rapor_no
            
            # Müşteri bilgisi
            musteri = "Bulunamadı"
            for pattern in musteri_patterns:
                if re.search(pattern, filename, re.IGNORECASE):
                    musteri = pattern
                    break
            values["musteri"] = musteri
            
            # Revizyon bilgisi
            revizyon_match = re.search(r'[vV](\d+)', filename)
            values["revizyon"] = f"v{revizyon_match.group(1)}" if revizyon_match else "v0"
        
        # Önemli değerler için pattern'ler - çok daha kapsamlı
        value_patterns = {
            # Proje adı/numarası için kapsamlı pattern'ler
            "proje_adi": [
                r"(?:Project\s*Name|Proje\s*Ad[ıi])\s*[:=]\s*([^\n\r]+)",
                r"(?:Project\s*No|Proje\s*No|Project\s*Number)\s*[:=]\s*([A-Z0-9.-]+)",
                r"(?:Report\s*Title|Rapor\s*Başl[ıi]ğ[ıi])\s*[:=]\s*([^\n\r]+)",
                r"(?:Document\s*Title|Belge\s*Başl[ıi]ğ[ıi])\s*[:=]\s*([^\n\r]+)",
                r"(LVD\s+[Öö]lç[üu]m[^,\n]*)",
                r"(Topraklama\s+S[üu]reklilik[^,\n]*)",
                r"(Grounding\s+Continuity[^,\n]*)",
                r"([A-Z][a-z]+\s*-\s*[A-Z][a-z]+.*?[Öö]lç[üu]m)",
                r"(E\d{2}\.\d{3}\s*-[^,\n]+)"
            ],
            
            # Rapor numarası için kapsamlı pattern'ler
            "rapor_numarasi": [
                r"(?:Report\s*No|Rapor\s*No|Report\s*Number)\s*[:=]\s*([A-Z0-9.-]+)",
                r"(?:Document\s*No|Belge\s*No)\s*[:=]\s*([A-Z0-9.-]+)",
                r"(E\d{2}\.\d{3})",
                r"(C\d{2}\.\d{3})",
                r"(T\d{2,3}[-.]?\d{3,4})",
                r"SM\s*(\d+)",
                r"MCC(\d+)",
                r"^\s*([A-Z]\d{2,3}[.-]\d{3,4})"
            ],
            
            # Ölçüm cihazı için çok kapsamlı pattern'ler
            "olcum_cihazi": [
                r"(?:Measuring\s*Instrument|Ölçüm\s*Cihaz[ıi]|Test\s*Equipment)\s*[:=]\s*([^\n\r]+)",
                r"(?:Multimeter|Multimetre|Ohmmeter|Ohmmetre)\s*[:=]?\s*([A-Z0-9\s.-]+)",
                r"(?:Instrument|Cihaz)\s*[:=]\s*([^\n\r]+)",
                r"(?:Equipment|Ekipman)\s*[:=]\s*([^\n\r]+)",
                r"(?:Device|Alet)\s*[:=]\s*([^\n\r]+)",
                r"(?:Tester|Test\s*Cihaz[ıi])\s*[:=]?\s*([A-Z0-9\s.-]+)",
                r"(Fluke\s*\d+[A-Z]*)",
                r"(Metrix\s*[A-Z0-9]+)",
                r"(Chauvin\s*Arnoux\s*[A-Z0-9]+)",
                r"(Megger\s*[A-Z0-9]+)",
                r"(Hioki\s*[A-Z0-9]+)",
                r"([A-Z][a-z]+\s*\d+[A-Z]*)",  # Genel marka model formatı
                r"(MΩ\s*metre|mΩ\s*metre|Loop\s*Tester|Continuity\s*Tester)"
            ],
            
            # Tesis/müşteri bilgisi
            "tesis_adi": [
                r"(?:Customer|Müşteri|Client)\s*[:=]\s*([^\n\r]+)",
                r"(?:Facility|Tesis|Plant|Factory)\s*[:=]\s*([^\n\r]+)",
                r"(?:Company|Firma|Corporation)\s*[:=]\s*([^\n\r]+)",
                r"(Toyota[^\n\r]*)",
                r"(DANONE[^\n\r]*)",
                r"(Ford[^\n\r]*)",
                r"(BOSCH[^\n\r]*)",
                r"(?:8X45|8X50|8X9J|9J73)\s*(?:R1|R2|R3)?\s*Hatt[ıi]",
                r"([A-Z][a-z]+\s+[A-Z][a-z]+\s+(?:Factory|Plant|Facility))"
            ],
            

            
            # Firma/personel bilgisi
            "firma_personel": [
                r"(?:Prepared\s*by|Hazırlayan|Consultant)\s*[:=]\s*([^\n\r]+)",
                r"(?:Performed\s*by|Ölçümü\s*Yapan)\s*[:=]\s*([^\n\r]+)",
                r"(?:Company|Firma)\s*[:=]\s*([^\n\r]+)",
                r"(?:Engineer|Mühendis)\s*[:=]\s*([^\n\r]+)",
                r"(PILZ[^\n\r]*)",
                r"([A-Z][a-z]+\s+[A-Z][a-z]+\s+(?:Engineering|Mühendislik))"
            ],
            
            # Tarih pattern'leri - çok kapsamlı
            "olcum_tarihi": [
                # Türkçe formatlar
                r"(?:Ölçüm\s*Tarihi|Test\s*Tarihi|Ölçüm\s*Yapıldığı\s*Tarih)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Ölçüm|Test).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"Tarih\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                
                # İngilizce formatlar
                r"(?:Measurement\s*Date|Test\s*Date|Date\s*of\s*(?:Test|Measurement))\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Measured\s*on|Tested\s*on)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Date|When)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                
                # Çeşitli formatlar
                r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})\s*(?:tarihinde|on|at|de)",
                r"(\d{4}[./\-]\d{1,2}[./\-]\d{1,2})",  # YYYY/MM/DD formatı
                r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})",  # İngilizce ay isimleri
                r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4})",  # Türkçe ay isimleri
                
                # Tablo içindeki tarihler
                r"(?:Measurement|Ölçüm).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",  # Genel tarih formatı
            ],
            
            "rapor_tarihi": [
                # Türkçe formatlar
                r"(?:Rapor\s*Tarihi|Belge\s*Tarihi|Hazırlanma\s*Tarihi)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Rapor|Belge).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Hazırlayan|Hazırlandı)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                
                # İngilizce formatlar  
                r"(?:Report\s*Date|Document\s*Date|Issue\s*Date|Prepared\s*Date)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Prepared\s*on|Issued\s*on|Created\s*on)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(?:Report|Document).*?(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                
                # Çeşitli formatlar
                r"(?:Date|Tarih)\s*[:=]\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(\d{4}[./\-]\d{1,2}[./\-]\d{1,2})",  # YYYY/MM/DD formatı
                r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})",  # İngilizce ay isimleri
                r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4})",  # Türkçe ay isimleri
                
                # Tablo başlığı veya footer'daki tarihler
                r"(?:Created|Issued|Published)\s*[:=]?\s*(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",
                r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{4})",  # Genel tarih formatı
            ]
        }
        
        # Metinden değerleri çıkar - her pattern listesi için
        for key, pattern_list in value_patterns.items():
            if key not in values:  # Dosya adından çıkarılmamışsa
                found_value = "Bulunamadı"
                
                # Pattern listesinde her pattern'i dene
                for pattern in pattern_list:
                    matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                    if matches:
                        if isinstance(matches[0], tuple):
                            # Tuple içindeki boş olmayan ilk değeri al
                            value = [m for m in matches[0] if m.strip()]
                            if value:
                                found_value = value[0].strip()
                                break
                        else:
                            found_value = matches[0].strip()
                            break
                
                values[key] = found_value
        
        # Ölçüm verilerini analiz et
        self.analyze_measurement_data(text, values)
        
        return values
    
    def analyze_measurement_data(self, text: str, values: Dict[str, Any]):
        """Ölçüm verilerini analiz et"""
        # RLO değerlerini topla - daha geniş pattern
        rlo_patterns = [
            r"(\d+[.,]?\d*)\s*(?:mΩ|mohm|ohm|Ω)",
            r"(\d+)\s*(?:4x[2-9](?:[.,]\d+)?|4x4)\s*(?:[2-9](?:[.,]\d+)?|4)\s*500",
            r"(\d+)\s*(?:mΩ|mohm|ohm|Ω)"
        ]
        
        rlo_values = []
        for pattern in rlo_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    # Virgülü noktaya çevir ve sayıya çevir
                    value_str = str(match).replace(',', '.')
                    rlo_values.append(float(value_str))
                except:
                    continue
        
        if rlo_values:
            values["rlo_min"] = f"{min(rlo_values):.1f} mΩ"
            values["rlo_max"] = f"{max(rlo_values):.1f} mΩ"
            values["rlo_ortalama"] = f"{sum(rlo_values)/len(rlo_values):.1f} mΩ"
        else:
            values["rlo_min"] = "Bulunamadı"
            values["rlo_max"] = "Bulunamadı"
            values["rlo_ortalama"] = "Bulunamadı"
        
        # Kesit bilgilerini analiz et - daha geniş pattern
        kesit_patterns = [
            r"4x4",
            r"4x2[.,]5", 
            r"4x6",
            r"4x10",
            r"Yük\s*İletken",
            r"Load\s*Conductor",
            r"PE\s*İletken",
            r"PE\s*Conductor"
        ]
        
        total_kesit_count = 0
        for pattern in kesit_patterns:
            count = len(re.findall(pattern, text, re.IGNORECASE))
            total_kesit_count += count
        
        values["toplam_olcum_nokta"] = total_kesit_count
        
        # Uygunluk durumlarını say
        uygun_pattern = r"UYGUNUYGUN"
        uygun_matches = re.findall(uygun_pattern, text)
        values["uygun_nokta_sayisi"] = len(uygun_matches)
        
        # Uygunsuz ölçümleri tespit et
        self.find_non_compliant_measurements(text, values)
        
        # Genel sonuç
        if len(uygun_matches) == values["toplam_olcum_nokta"] and values["toplam_olcum_nokta"] > 0:
            values["genel_sonuc"] = "TÜM NOKTALAR UYGUN"
        else:
            values["genel_sonuc"] = f"{values['toplam_olcum_nokta'] - len(uygun_matches)} NOKTA UYGUNSUZ"
        
        # Hat/bölge bilgileri
        hat_pattern = r"(8X45|8X50|8X9J|9J73|8X52|8X60|8X62|8X70)\s*(?:R[1-9])?\s*Hatt[ıi]"
        hat_matches = re.findall(hat_pattern, text, re.IGNORECASE)
        if hat_matches:
            unique_hats = list(set(hat_matches))
            values["makine_hatlari"] = ", ".join(unique_hats)
        else:
            values["makine_hatlari"] = "Bulunamadı"
    
    def find_non_compliant_measurements(self, text: str, values: Dict[str, Any]):
        """Uygunsuz ölçümleri tespit et"""
        # 500 mΩ'dan büyük değerleri ve D.Y. değerlerini bul
        lines = text.split('\n')
        non_compliant = []
        
        for i, line in enumerate(lines):
            # Sıra numarası kontrolü
            sira_match = re.search(r'(\d+)\s', line)
            if sira_match:
                sira = sira_match.group(1)
                
                # Yüksek RLO değeri kontrolü (>500 mΩ) - daha geniş pattern
                high_rlo_patterns = [
                    r'(\d{3,4})\s*(?:4x[2-9](?:[.,]\d+)?|4x4)\s*(?:[2-9](?:[.,]\d+)?|4)\s*500(\d+)\s*mΩ\s*<\s*500\s*mΩ',
                    r'(\d{3,4})\s*(?:mΩ|mohm|ohm|Ω)',
                    r'(\d{3,4})[.,]?\d*\s*(?:mΩ|mohm|ohm|Ω)'
                ]
                
                for pattern in high_rlo_patterns:
                    high_rlo_match = re.search(pattern, line, re.IGNORECASE)
                    if high_rlo_match:
                        try:
                            rlo_value = float(str(high_rlo_match.group(1)).replace(',', '.'))
                            if rlo_value > 500:
                                # Hat ve ekipman bilgisi - daha geniş pattern
                                hat_patterns = [
                                    r'(8X\d+R?\d*)\s*(?:Hatt[ıi]|Line|Zone)?\s*(.*?)(?:\s+\d+)',
                                    r'(8X\d+R?\d*)\s*(.*?)(?:\s+\d+)',
                                    r'(Line\s*\d+|Zone\s*\d+)\s*(.*?)(?:\s+\d+)'
                                ]
                                
                                for hat_pattern in hat_patterns:
                                    hat_match = re.search(hat_pattern, line, re.IGNORECASE)
                                    if hat_match:
                                        hat = hat_match.group(1)
                                        ekipman = hat_match.group(2).strip()
                                        non_compliant.append({
                                            'sira': sira,
                                            'rlo': f"{rlo_value:.1f} mΩ",
                                            'hat': hat,
                                            'ekipman': ekipman,
                                            'durum': 'Yüksek Direnç'
                                        })
                                        break
                                break
                        except:
                            continue
                
                # D.Y. (Değer Yok) kontrolü - daha geniş pattern
                if '*D.Y' in line or 'D.Y' in line or 'N/A' in line or 'N/A' in line:
                    hat_patterns = [
                        r'(8X\d+R?\d*)\s*(?:Hatt[ıi]|Line|Zone)?\s*(.*?)(?:\s+|$)',
                        r'(8X\d+R?\d*)\s*(.*?)(?:\s+|$)',
                        r'(Line\s*\d+|Zone\s*\d+)\s*(.*?)(?:\s+|$)'
                    ]
                    
                    for hat_pattern in hat_patterns:
                        hat_match = re.search(hat_pattern, line, re.IGNORECASE)
                        if hat_match:
                            hat = hat_match.group(1)
                            ekipman = hat_match.group(2).strip()
                            non_compliant.append({
                                'sira': sira,
                                'rlo': 'D.Y.',
                                'hat': hat,
                                'ekipman': ekipman,
                                'durum': 'Ölçüm Yapılamadı'
                            })
                            break
        
        values["uygunsuz_olcumler"] = non_compliant
    
    def calculate_scores(self, analysis_results: Dict[str, Dict[str, GroundingAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesaplama"""
        category_scores = {}
        total_score = 0
        total_max_score = 100
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            # Kategori puanını ağırlığa göre normalize et
            normalized_score = (category_earned / category_possible * category_max) if category_possible > 0 else 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round((category_earned / category_possible * 100), 2) if category_possible > 0 else 0
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "total_max_score": total_max_score,
            "overall_percentage": round((total_score / total_max_score * 100), 2)
        }
    
    def generate_detailed_report(self, file_path: str) -> Dict[str, Any]:
        """Detaylı rapor oluşturma"""
        logger.info("Topraklama Süreklilik rapor analizi başlatılıyor...")
        
        # Dosyadan metin çıkar ve dil bilgisi al
        text, detected_language = self.get_file_text(file_path)
        if not text:
            return {"error": "Dosya okunamadı"}
        
        # Dil bilgisini logla
        language_name = self.get_language_name(detected_language)
        logger.info(f"📖 Belge dili: {language_name}")
        if detected_language != 'tr':
            logger.info("🔄 Çeviri işlemi tamamlandı")
        
        # Tarih geçerliliği kontrolü (1 yıl kuralı)
        date_valid, olcum_tarihi, rapor_tarihi, date_message = self.check_date_validity(text, file_path)
        
        # Spesifik değerleri çıkar
        extracted_values = self.extract_specific_values(text, file_path)
        
        # Dil bilgisini extracted_values'a ekle
        extracted_values['detected_language'] = detected_language
        extracted_values['language_name'] = language_name
        
        # Her kategori için analiz yap
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        # Puanları hesapla
        scores = self.calculate_scores(analysis_results)
        
        # Final karar: Tarih geçersizse puan ne olursa olsun FAILED
        final_status = "PASSED"
        if not date_valid:
            final_status = "FAILED"
            fail_reason = "Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla"
        elif scores["overall_percentage"] < 70:
            final_status = "FAILED"
            fail_reason = f"Toplam puan yetersiz (%{scores['overall_percentage']:.1f} < 70)"
        else:
            fail_reason = None
        
        # Öneriler oluştur
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid)
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgileri": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1]
            },
            "tarih_gecerliligi": {
                "gecerli": date_valid,
                "olcum_tarihi": olcum_tarihi,
                "rapor_tarihi": rapor_tarihi,
                "mesaj": date_message
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["overall_percentage"],
                "final_durum": final_status,
                "tarih_durumu": "GEÇERLİ" if date_valid else "GEÇERSİZ",
                "gecme_durumu": "PASSED" if final_status == "PASSED" else "FAILED",
                "fail_nedeni": fail_reason
            }
        }
        
        return report
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool) -> List[str]:
        """Öneriler oluşturma"""
        recommendations = []
        
        # Tarih kontrolü öncelikli
        if not date_valid:
            recommendations.append("🚨 KRİTİK: Ölçüm tarihi ile rapor tarihi arasındaki fark 1 yıldan fazla - RAPOR GEÇERSİZ")
            recommendations.append("- Yeni ölçüm yapılması gereklidir")
            recommendations.append("- Ölçüm tarihi rapor tarihinden en fazla 1 yıl önce olmalıdır")
        
        # Kategori bazlı öneriler
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 50:
                recommendations.append(f"❌ {category} bölümü yetersiz (%{category_score:.1f})")
                
                # Eksik kriterler
                missing_criteria = [name for name, result in results.items() if not result.found]
                if missing_criteria:
                    recommendations.append(f"  Eksik kriterler: {', '.join(missing_criteria)}")
                
                # Kategori özel öneriler
                if category == "Genel Rapor Bilgileri":
                    recommendations.append("  - Proje adı ve numarası eksiksiz belirtilmelidir")
                    recommendations.append("  - Ölçüm ve rapor tarihleri açıkça belirtilmelidir")
                    recommendations.append("  - Rapor numarası ve revizyon bilgisi eklenmeli")
                
                elif category == "Ölçüm Metodu ve Standart Referansları":
                    recommendations.append("  - Ölçüm cihazı marka/model bilgileri eklenmeli")
                    recommendations.append("  - Kalibrasyon sertifikası bilgileri verilmeli")
                    recommendations.append("  - EN 60204-1 Tablo 10 referansı yapılmalı")
                
                elif category == "Ölçüm Sonuç Tablosu":
                    recommendations.append("  - Tüm ölçüm noktaları için RLO değerleri belirtilmeli")
                    recommendations.append("  - Yük ve PE iletken kesitleri girilmeli")
                    recommendations.append("  - EN 60204 Tablo 10 referans değerleri eklenmeli")
                    recommendations.append("  - Uygunluk durumu her nokta için belirtilmeli")
                
                elif category == "Uygunluk Değerlendirmesi":
                    recommendations.append("⚠️ Uygunsuz noktalar için teknik açıklama ekleyin")
                    recommendations.append("📊 Toplam ölçüm sayısı ve uygunluk oranını belirtin")
                    recommendations.append("🔍 500 mΩ limit değeri aşımlarını işaretleyin")
                
                elif category == "Görsel ve Teknik Dökümantasyon":
                    recommendations.append("  - Ölçüm yapılan alan fotoğrafları eklenmeli")
                    recommendations.append("  - Ölçüm cihazı ve bağlantı fotoğrafları çekilmeli")
                    recommendations.append("  - Ölçüm noktalarının kroki/şeması hazırlanmalı")
                
                elif category == "Sonuç ve Öneriler":
                    recommendations.append("  - Genel uygunluk sonucu açıkça belirtilmeli")
                    recommendations.append("  - Standartlara atıf yapılmalı")
                    recommendations.append("  - İyileştirme önerileri detaylandırılmalı")
                    recommendations.append("  - Tekrar ölçüm periyodu önerilmeli")
            
            elif category_score < 80:
                recommendations.append(f"⚠️ {category} bölümü geliştirilmeli (%{category_score:.1f})")
            
            else:
                recommendations.append(f"✅ {category} bölümü yeterli (%{category_score:.1f})")
        
        # Genel öneriler
        if scores["overall_percentage"] < 70:
            recommendations.append("\n🚨 GENEL ÖNERİLER:")
            recommendations.append("- Rapor EN 60204-1 standardına tam uyumlu hale getirilmelidir")
            recommendations.append("- IEC 60364 standart referansları eklenmeli")
            recommendations.append("- Eksik bilgiler tamamlanmalıdır")
            recommendations.append("- Ölçüm sonuçları tablo formatında düzenlenmeli")
        
        # Başarılı durumda
        if scores["overall_percentage"] >= 70 and date_valid:
            recommendations.append("\n✅ RAPOR BAŞARILI")
            recommendations.append("- Tüm gerekli kriterler sağlanmıştır")
            recommendations.append("- Rapor standarltara uygun olarak hazırlanmıştır")
        
        return recommendations
    
    def save_report_to_excel(self, report: Dict, output_path: str):
        """Raporu Excel'e kaydetme"""
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Özet sayfa
            ozet_data = {
                'Kriter': ['Toplam Puan', 'Yüzde', 'Final Durum', 'Tarih Durumu', 'Geçme Durumu'],
                'Değer': [
                    report['ozet']['toplam_puan'],
                    f"%{report['ozet']['yuzde']}",
                    report['ozet']['final_durum'],
                    report['ozet']['tarih_durumu'],
                    report['ozet']['gecme_durumu']
                ]
            }
            if report['ozet']['fail_nedeni']:
                ozet_data['Kriter'].append('Başarısızlık Nedeni')
                ozet_data['Değer'].append(report['ozet']['fail_nedeni'])
            
            pd.DataFrame(ozet_data).to_excel(writer, sheet_name='Özet', index=False)
            
            # Çıkarılan değerler
            values_data = []
            for key, value in report['cikarilan_degerler'].items():
                values_data.append({'Kriter': key, 'Değer': value})
            pd.DataFrame(values_data).to_excel(writer, sheet_name='Çıkarılan Değerler', index=False)
            
            # Kategori detayları
            for category, results in report['kategori_analizleri'].items():
                category_data = []
                for criterion, result in results.items():
                    category_data.append({
                        'Kriter': criterion,
                        'Bulundu': result.found,
                        'İçerik': result.content,
                        'Puan': result.score,
                        'Max Puan': result.max_score
                    })
                
                sheet_name = category[:31]  # Excel sheet name limit
                pd.DataFrame(category_data).to_excel(writer, sheet_name=sheet_name, index=False)
        
        logger.info(f"Rapor Excel dosyası kaydedildi: {output_path}")
    
    def save_report_to_json(self, report: Dict, output_path: str):
        """Raporu JSON'a kaydetme"""
        # GroundingAnalysisResult objelerini dict'e çevir
        json_report = {}
        for key, value in report.items():
            if key == 'kategori_analizleri':
                json_report[key] = {}
                for category, results in value.items():
                    json_report[key][category] = {}
                    for criterion, result in results.items():
                        json_report[key][category][criterion] = asdict(result)
            else:
                json_report[key] = value
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_report, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Rapor JSON dosyası kaydedildi: {output_path}")

def main():
    """Ana fonksiyon"""
    analyzer = GroundingContinuityReportAnalyzer()
    
    # Dosya yolu - Proje root'undaki belgeyi analiz et
    file_path = "E21.207 - Toyota - Chifong LVD Ölçüm .pdf"
    
    # Dosyanın varlığını kontrol et
    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        print("Mevcut dosyalar:")
        for file in os.listdir("."):
            if file.endswith(('.pdf', '.docx', '.xlsx')):
                print(f"  - {file}")
        return
    
    print("⚡ Topraklama Süreklilik Rapor Analizi Başlatılıyor...")
    print("=" * 60)
    
    # Analizi çalıştır
    report = analyzer.generate_detailed_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    print("\n📊 ANALİZ SONUÇLARI")
    print("=" * 60)
    
    print(f"📅 Analiz Tarihi: {report['analiz_tarihi']}")
    print(f"🔍 Tespit Edilen Dil: {report['cikarilan_degerler'].get('language_name', 'Bilinmiyor')}")
    print(f"📋 Toplam Puan: {report['ozet']['toplam_puan']}/100")
    print(f"📈 Yüzde: %{report['ozet']['yuzde']:.1f}")
    print(f"🎯 Durum: {report['ozet']['final_durum']}")
    print(f"📄 Rapor Tipi: LVD Ölçüm Raporu")
    
    print(f"\n📅 TARİH GEÇERLİLİĞİ")
    print("-" * 40)
    print(f"Ölçüm Tarihi: {report['tarih_gecerliligi']['olcum_tarihi']}")
    print(f"Rapor Tarihi: {report['tarih_gecerliligi']['rapor_tarihi']}")
    print(f"Geçerlilik: {report['tarih_gecerliligi']['mesaj']}")
    
    print("\n📋 ÖNEMLİ ÇIKARILAN DEĞERLER")
    print("-" * 40)
    important_values = {
        "proje_no": "Proje No",
        "rapor_numarasi": "Rapor Numarası", 
        "musteri": "Müşteri",
        "proje_adi": "Proje Adı",
        "tesis_adi": "Tesis/Hat",
        "olcum_cihazi": "Ölçüm Cihazı",
        "firma_personel": "Hazırlayan Firma",
        "toplam_olcum_nokta": "Toplam Ölçüm Noktası",
        "uygun_nokta_sayisi": "Uygun Nokta Sayısı",
        "genel_sonuc": "Genel Sonuç"
    }
    
    for key, display_name in important_values.items():
        if key in report['cikarilan_degerler']:
            print(f"{display_name}: {report['cikarilan_degerler'][key]}")
    
    print("\n📊 KATEGORİ PUANLARI VE DETAYLAR")
    print("=" * 60)
    
    # Her kategori için detaylı analiz
    categories = [
        ("Genel Rapor Bilgileri", "1"),
        ("Ölçüm Metodu ve Standart Referansları", "2"), 
        ("Ölçüm Sonuç Tablosu", "3"),
        ("Uygunluk Değerlendirmesi", "4"),
        ("Görsel ve Teknik Dökümantasyon", "5"),
        ("Sonuç ve Öneriler", "6")
    ]
    
    for category, num in categories:
        if category in report['puanlama']['category_scores']:
            score_data = report['puanlama']['category_scores'][category]
            percentage = score_data['percentage']
            print(f"\n🔍 {category}: {score_data['normalized']:.1f}/{score_data['max_weight']} (%{percentage:.1f})")
            print("-" * 50)
            
            # Bu kategorinin analiz sonuçlarını göster
            if category in report['kategori_analizleri']:
                category_analysis = report['kategori_analizleri'][category]
                for criterion_name, criterion_result in category_analysis.items():
                    criterion_display = criterion_name.replace('_', ' ').title()
                    if hasattr(criterion_result, 'found') and criterion_result.found:
                        print(f"  ✅ {criterion_display}: {criterion_result.score}/{criterion_result.max_score} puan")
                    else:
                        print(f"  ❌ {criterion_display}: 0/{criterion_result.max_score} puan - BULUNAMADI")
    
    print("\n" + "=" * 60)
    
    print("\n💡 ÖNERİLER VE DEĞERLENDİRME")
    print("-" * 40)
    for recommendation in report['oneriler']:
        print(recommendation)
    
    print("\n📋 GENEL DEĞERLENDİRME")
    print("=" * 60)
    
    if report['ozet']['final_durum'] == "PASSED":
        print("✅ SONUÇ: GEÇERLİ")
        print(f"🌟 Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Topraklama Süreklilik raporu genel olarak yeterli kriterleri sağlamaktadır.")
    else:
        print("❌ SONUÇ: GEÇERSİZ")
        print(f"⚠️ Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Topraklama Süreklilik raporu minimum gereklilikleri sağlamamaktadır.")
        
        # Başarısızlık nedeni varsa yazdır
        if report['ozet']['fail_nedeni']:
            print(f"🚨 Başarısızlık Nedeni: {report['ozet']['fail_nedeni']}")
        
        print("\n⚠️ EKSİK GEREKLİLİKLER:")
        for category, results in report['kategori_analizleri'].items():
            missing_items = []
            for criterion, result in results.items():
                if not result.found:
                    missing_items.append(criterion)
            
            if missing_items:
                print(f"\n🔍 {category}:")
                for item in missing_items[:3]:  # İlk 3 eksik item'ı göster
                    readable_name = item.replace('_', ' ').title()
                    print(f"   ❌ {readable_name}")
        
        print("\n📌 YAPILMASI GEREKENLER:")
        print("1. Eksik belgelendirmeleri tamamlayın")
        print("2. Ölçüm cihazı ve kalibrasyon bilgilerini ekleyin")
        print("3. Ölçüm sonuç tablolarını detaylandırın")
        print("4. Uygunluk değerlendirmelerini güçlendirin")
        print("5. Görsel dokümantasyonu artırın")
        print("6. Standart referanslarını ekleyin")
        
        # Uygunsuz ölçümler varsa ekstra bilgi
        uygunsuz_olcumler = report['cikarilan_degerler'].get('uygunsuz_olcumler', [])
        if uygunsuz_olcumler:
            print("\n🚨 UYGUNSUZ ÖLÇÜMLER:")
            for measurement in uygunsuz_olcumler[:5]:  # İlk 5 uygunsuz ölçümü göster
                if measurement['durum'] == 'Yüksek Direnç':
                    print(f"   ⚠️ Sıra {measurement['sira']}: {measurement['rlo']} > 500 mΩ")
                else:
                    print(f"   ⚠️ Sıra {measurement['sira']}: Ölçüm yapılamadı (*D.Y.)")
            
            if len(uygunsuz_olcumler) > 5:
                print(f"   ... ve {len(uygunsuz_olcumler) - 5} uygunsuz ölçüm daha")

if __name__ == "__main__":
    main()
